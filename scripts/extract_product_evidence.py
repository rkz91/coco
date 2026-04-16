#!/usr/bin/env python3
"""
Product Evidence Extraction Pipeline
=====================================
Extracts text from internal docs for 25 Cross Risk products and outputs
structured evidence JSONs to ~/.coco/knowledge/product_evidence/.

Usage:
    python scripts/extract_product_evidence.py

Sources:
    - Cross-Risk-Overview.html (React SPA presentation)
    - PRD .docx / .md / .html files
    - Architecture diagrams (4+1 UML PNGs, PDFs)
    - 3PI V2 Architecture & PRD HTML files

Output:
    ~/.coco/knowledge/product_evidence/<slug>.json   (one per product)
    ~/.coco/knowledge/product_evidence/extraction_manifest.json
"""

from __future__ import annotations

import hashlib
import json
import logging
import os
import re
import sys
import traceback
from dataclasses import asdict, dataclass, field
from pathlib import Path
from typing import Optional

# ---------------------------------------------------------------------------
# Logging
# ---------------------------------------------------------------------------
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    datefmt="%H:%M:%S",
)
log = logging.getLogger("extract_evidence")

# ---------------------------------------------------------------------------
# Output directory
# ---------------------------------------------------------------------------
OUTPUT_DIR = Path.home() / ".coco" / "knowledge" / "product_evidence"

# ---------------------------------------------------------------------------
# Base search paths (resolved at runtime)
# ---------------------------------------------------------------------------
HOME = Path.home()
DOWNLOADS = HOME / "Downloads"
OLD_DOWNLOADS = DOWNLOADS / "Old Downloads"
DOCUMENTS = HOME / "Documents"
CONFLUENCE_PRD_DIR = DOCUMENTS / "Integration with Confluence" / "PRD"
ONEDRIVE_CROSS_RISK = (
    HOME
    / "Library/CloudStorage/OneDrive-McKinsey&Company"
    / "(INT)-Risk Products and Services - Cross Risk"
    / "I - Cross Risk"
)
ONEDRIVE_INITIATIVES = (
    HOME
    / "Library/CloudStorage/OneDrive-McKinsey&Company"
    / "(INT)-Risk Products and Services - Initiatives"
)
ONEDRIVE_PRDS = ONEDRIVE_CROSS_RISK / "Cross Risk Internal" / "PRD's"
ARCH_UML_DIR = ONEDRIVE_CROSS_RISK / "Cross Risk Internal" / "Architecture Diagrams - 4+1 UML"
THREE_PI_V2_DIR = ONEDRIVE_CROSS_RISK / "Optimize" / "3PI V2"
CFC_DIR = ONEDRIVE_CROSS_RISK / "Optimize" / "1024 - CFC Third Party Management"
AUDITBOARD_DIR = ONEDRIVE_CROSS_RISK / "Audit Board"
REG_COMP_DIR = ONEDRIVE_CROSS_RISK / "E&C" / "Regulatory Compliance"
PROJECT_ROOT = Path(__file__).resolve().parent.parent

# ---------------------------------------------------------------------------
# Cross-Risk-Overview.html search paths (ordered by preference)
# ---------------------------------------------------------------------------
CROSS_RISK_OVERVIEW_CANDIDATES = [
    OLD_DOWNLOADS / "Cross-Risk-Overview.html",
    DOWNLOADS / "Cross-Risk-Overview.html",
    PROJECT_ROOT / "Cross-Risk-Overview.html",
]


# ---------------------------------------------------------------------------
# Data structures
# ---------------------------------------------------------------------------
@dataclass
class EvidenceSource:
    type: str  # presentation, prd, architecture, reference
    name: str
    content: str
    source_path: str = ""
    char_count: int = 0

    def __post_init__(self):
        self.char_count = len(self.content)


@dataclass
class ProductEvidence:
    product: str
    slug: str
    team: str
    program: str
    project_id: str = ""
    sources: list[EvidenceSource] = field(default_factory=list)
    evidence_tier: str = "thin"  # rich | standard | thin
    source_count: int = 0
    source_hash: str = ""
    errors: list[str] = field(default_factory=list)

    def compute_tier(self):
        """Determine evidence richness: rich (PRD+arch), standard (overview+PRD), thin (overview only)."""
        types = {s.type for s in self.sources}
        total_chars = sum(s.char_count for s in self.sources)
        has_prd = "prd" in types
        has_arch = "architecture" in types
        has_presentation = "presentation" in types

        if has_prd and has_arch and total_chars > 5000:
            self.evidence_tier = "rich"
        elif (has_prd or has_arch) and total_chars > 2000:
            self.evidence_tier = "standard"
        else:
            self.evidence_tier = "thin"

        self.source_count = len(self.sources)

    def compute_hash(self):
        """SHA-256 of all source content for staleness detection."""
        h = hashlib.sha256()
        for s in sorted(self.sources, key=lambda x: x.name):
            h.update(s.content.encode("utf-8", errors="replace"))
        self.source_hash = h.hexdigest()


# ---------------------------------------------------------------------------
# 25 Products — master mapping
# ---------------------------------------------------------------------------
PRODUCTS: list[dict] = [
    # --- Anti-Corruption ---
    {
        "product": "TPI Tracker",
        "slug": "tpi-tracker",
        "team": "Anti-Corruption",
        "program": "E&C",
        "project_id": "19986",
        "overview_program": "Anti-Corruption",
        "overview_keywords": ["TPI Tracker", "TPI tracker", "tracker", "third-party intermediary"],
        "prds": [
            ("Product Requirement Document - TPI Tracker.docx", "docx"),
            ("Product Requirement Document \u2013 TPI Risk Ranking.docx", "docx"),
        ],
        "architecture": ["TPI Tracker"],  # subdirs under ARCH_UML_DIR
        "arch_pdfs": ["TPI Tracker Architecture.pdf"],
    },
    {
        "product": "OneTrust TPI-TPDD",
        "slug": "onetrust-tpi-tpdd",
        "team": "Anti-Corruption",
        "program": "E&C",
        "project_id": "19987",
        "overview_program": "Anti-Corruption",
        "overview_keywords": ["TPI-TPDD", "TPDD", "Third Party Due Diligence", "OneTrust"],
        "prds": [],
        "architecture": [],
    },
    {
        "product": "TPI Central Data Repo",
        "slug": "tpi-central-data-repo",
        "team": "Anti-Corruption",
        "program": "E&C",
        "project_id": "",
        "overview_program": "Anti-Corruption",
        "overview_keywords": ["Central Data", "data repo", "TPI Central"],
        "prds": [],
        "architecture": [],
    },
    {
        "product": "GL Monitoring",
        "slug": "gl-monitoring",
        "team": "Anti-Corruption",
        "program": "E&C",
        "project_id": "",
        "overview_program": "Anti-Corruption",
        "overview_keywords": ["GL Monitoring", "GL Screening", "Gift", "Ledger", "General Ledger"],
        "prds": [
            ("Product Requirement Document \u2013 GL Screening.docx", "docx"),
        ],
        "architecture": [],
    },
    # --- Privacy ---
    {
        "product": "OneTrust Cookie Compliance",
        "slug": "onetrust-cookie-compliance",
        "team": "Privacy",
        "program": "E&C",
        "project_id": "",
        "overview_program": "Privacy",
        "overview_keywords": ["Cookie Compliance", "cookie", "web cookie"],
        "prds": [
            ("Product Requirement Document \u2013 Web Cookie Compliance.docx", "docx"),
        ],
        "architecture": [],
    },
    {
        "product": "OneTrust Privacy Rights",
        "slug": "onetrust-privacy-rights",
        "team": "Privacy",
        "program": "E&C",
        "project_id": "",
        "overview_program": "Privacy",
        "overview_keywords": ["Privacy Rights", "DSR", "Data Subject Rights", "data subject request"],
        "prds": [],
        "architecture": [],
    },
    {
        "product": "OneTrust Mobile App Consent",
        "slug": "onetrust-mobile-app-consent",
        "team": "Privacy",
        "program": "E&C",
        "project_id": "",
        "overview_program": "Privacy",
        "overview_keywords": ["Mobile App Consent", "mobile consent", "app consent"],
        "prds": [],
        "architecture": [],
    },
    {
        "product": "OneTrust Data Mapping",
        "slug": "onetrust-data-mapping",
        "team": "Privacy",
        "program": "E&C",
        "project_id": "",
        "overview_program": "Privacy",
        "overview_keywords": ["Data Mapping", "Article 30", "data inventory", "ROPA"],
        "prds": [
            ("Product Requirement Document \u2013 Privacy Data Mapping Article 30 final docx.docx", "docx"),
        ],
        "architecture": [],
    },
    {
        "product": "OneTrust Consent Rate Optimization",
        "slug": "onetrust-consent-rate-opt",
        "team": "Privacy",
        "program": "E&C",
        "project_id": "",
        "overview_program": "Privacy",
        "overview_keywords": ["Consent Rate", "consent optimization", "opt-in rate"],
        "prds": [],
        "architecture": [],
    },
    {
        "product": "OneTrust Universal Cookie",
        "slug": "onetrust-universal-cookie",
        "team": "Privacy",
        "program": "E&C",
        "project_id": "",
        "overview_program": "Privacy",
        "overview_keywords": ["Universal Cookie", "universal consent"],
        "prds": [],
        "architecture": [],
    },
    {
        "product": "OneTrust Data Guidance",
        "slug": "onetrust-data-guidance",
        "team": "Privacy",
        "program": "E&C",
        "project_id": "",
        "overview_program": "Privacy",
        "overview_keywords": ["Data Guidance", "DataGuidance", "regulatory research"],
        "prds": [],
        "architecture": [],
    },
    {
        "product": "OneTrust PIA DPIA",
        "slug": "onetrust-pia-dpia",
        "team": "Privacy",
        "program": "E&C",
        "project_id": "",
        "overview_program": "Privacy",
        "overview_keywords": ["PIA", "DPIA", "Privacy Impact", "Data Protection Impact"],
        "prds": [],
        "architecture": [],
    },
    # --- Regulatory Compliance ---
    {
        "product": "Bridger",
        "slug": "bridger",
        "team": "Regulatory Compliance",
        "program": "E&C",
        "project_id": "17654",
        "overview_program": "Regulatory Compliance",
        "overview_keywords": ["Bridger", "LexisNexis", "screening", "sanctions screening"],
        "prds": [
            ("Product Requirement Document \u2013 Bridger and Reg COE Portal Improvement.docx", "docx"),
            ("Product Requirement Document \u2013 Counter party data extract.docx", "docx"),
        ],
        "architecture": [],
    },
    {
        "product": "Reg Compliance Screening Portal",
        "slug": "reg-compliance-screening",
        "team": "Regulatory Compliance",
        "program": "E&C",
        "project_id": "17654",
        "overview_program": "Regulatory Compliance",
        "overview_keywords": ["Screening Portal", "Reg CoE", "Request Portal", "sanctions portal"],
        "prds": [
            ("Reg CoE Requirements Document.docx", "docx"),
            ("Product Requirement Document \u2013 Bridger and Reg COE Portal Improvement.docx", "docx"),
        ],
        "architecture": ["Request Portal"],
    },
    {
        "product": "AI Case Operator",
        "slug": "ai-case-operator",
        "team": "Regulatory Compliance",
        "program": "E&C",
        "project_id": "",
        "overview_program": "Regulatory Compliance",
        "overview_keywords": ["AI Case Operator", "case operator", "AI-powered", "case review", "AI agent"],
        "prds": [],
        "architecture": [],
    },
    {
        "product": "Moodys Orbis",
        "slug": "moodys-orbis",
        "team": "Regulatory Compliance",
        "program": "E&C",
        "project_id": "19522",
        "overview_program": "Regulatory Compliance",
        "overview_keywords": ["Orbis", "Moody", "company data", "beneficial ownership"],
        "prds": [],
        "architecture": [],
        "arch_pdfs": ["Orbis_Architecture_Diagram.pdf"],
    },
    {
        "product": "Kharon",
        "slug": "kharon",
        "team": "Regulatory Compliance",
        "program": "E&C",
        "project_id": "",
        "overview_program": "Regulatory Compliance",
        "overview_keywords": ["Kharon", "sanctions intelligence", "risk intelligence"],
        "prds": [],
        "architecture": [],
    },
    # --- External tools (internet research needed) ---
    {
        "product": "MLex",
        "slug": "mlex",
        "team": "Regulatory Compliance",
        "program": "E&C",
        "project_id": "",
        "overview_program": None,
        "overview_keywords": ["MLex"],
        "prds": [],
        "architecture": [],
    },
    {
        "product": "TransparINT",
        "slug": "transparint",
        "team": "Regulatory Compliance",
        "program": "E&C",
        "project_id": "",
        "overview_program": None,
        "overview_keywords": ["TransparINT"],
        "prds": [],
        "architecture": [],
    },
    {
        "product": "Navex",
        "slug": "navex",
        "team": "Regulatory Compliance",
        "program": "E&C",
        "project_id": "",
        "overview_program": None,
        "overview_keywords": ["Navex", "NAVEX"],
        "prds": [],
        "architecture": [],
    },
    {
        "product": "MyComplianceOffice",
        "slug": "mycomplianceoffice",
        "team": "Regulatory Compliance",
        "program": "E&C",
        "project_id": "",
        "overview_program": None,
        "overview_keywords": ["MyComplianceOffice", "MCO"],
        "prds": [],
        "architecture": [],
    },
    # --- AuditBoard ---
    {
        "product": "AuditBoard (AB1)",
        "slug": "auditboard-ab1",
        "team": "AuditBoard",
        "program": "Risk & Compliance",
        "project_id": "",
        "overview_program": "AuditBoard",
        "overview_keywords": ["AuditBoard", "AB1", "OpsAudit", "SOXHUB", "RiskOversight", "CrossComply"],
        "prds": [],
        "architecture": [],
    },
    {
        "product": "AuditBoard (AB2)",
        "slug": "auditboard-ab2",
        "team": "AuditBoard",
        "program": "Risk & Compliance",
        "project_id": "",
        "overview_program": "AuditBoard",
        "overview_keywords": ["AB2", "Tax Control", "TCF", "Tax Control Framework"],
        "prds": [
            ("AB2-PRD-Tax-Control-Framework.docx", "docx"),
        ],
        "architecture": [],
    },
    # --- Optimize ---
    {
        "product": "TP Inventory / 3PI",
        "slug": "tp-inventory",
        "team": "Optimize",
        "program": "Optimize",
        "project_id": "20100",
        "overview_program": "Optimize",
        "overview_keywords": ["TP Inventory", "3PI", "Third Party Inventory", "third-party inventory", "transfer pricing"],
        "prds": [
            ("Product Requirement Document TP Inventory.docx", "docx"),
            ("PRD - TP Inventory Natural Language Query.md", "md"),
            ("Product Requirement Document - TP Inventory Leadership.docx", "docx"),
        ],
        "architecture": ["TP Inventory"],
        "extra_arch": ["3PI V2"],  # Architecture.html + PRD-Phase2.html
    },
    {
        "product": "CFC Third Party Management",
        "slug": "cfc-tpm",
        "team": "Optimize",
        "program": "Optimize",
        "project_id": "1024",
        "overview_program": None,
        "overview_keywords": ["CFC", "Third Party Management", "CFC TPM"],
        "prds": [
            ("1024-CFC-Third-Party-Management-PRD.md", "md"),
        ],
        "architecture": [],
    },
]


# ---------------------------------------------------------------------------
# PRD file search resolution
# ---------------------------------------------------------------------------
# Map PRD filenames to known absolute paths (discovered during development)
PRD_SEARCH_DIRS: list[Path] = [
    CONFLUENCE_PRD_DIR,
    ONEDRIVE_PRDS,
]

# Special PRD locations that don't follow the standard directory pattern
SPECIAL_PRD_PATHS: dict[str, Path] = {
    "Reg CoE Requirements Document.docx": REG_COMP_DIR / "Reg CoE Requirements Document.docx",
    "AB2-PRD-Tax-Control-Framework.docx": AUDITBOARD_DIR / "AB2-Tax-Control-Framework" / "PRD" / "AB2-PRD-Tax-Control-Framework.docx",
    "Product Requirement Document TP Inventory.docx": (
        ONEDRIVE_CROSS_RISK / "Optimize" / "TP Inventory" / "Product Requirement Document TP Inventory.docx"
    ),
    "1024-CFC-Third-Party-Management-PRD.md": CFC_DIR / "1024-CFC-Third-Party-Management-PRD.md",
    "PRD - TP Inventory Natural Language Query.md": ONEDRIVE_PRDS / "PRD - TP Inventory Natural Language Query.md",
    "Product Requirement Document - TP Inventory Leadership.docx": (
        CONFLUENCE_PRD_DIR / "Product Requirement Document - TP Inventory Leadership.docx"
    ),
}

# Architecture PDF search paths
ARCH_PDF_PATHS: dict[str, Path] = {
    "TPI Tracker Architecture.pdf": OLD_DOWNLOADS / "Work - McKinsey" / "TPI & Vendor Data" / "TPI Tracker Architecture.pdf",
    "Orbis_Architecture_Diagram.pdf": REG_COMP_DIR / "Orbis" / "Orbis_Architecture_Diagram.pdf",
}

# 3PI V2 extra architecture files
THREE_PI_V2_DOCS: dict[str, Path] = {
    "3PI V2 Architecture": THREE_PI_V2_DIR / "docs" / "Architecture.html",
    "3PI V2 PRD-Phase2": THREE_PI_V2_DIR / "docs" / "PRD-Phase2.html",
    "3PI V2 Architecture Decision": THREE_PI_V2_DIR / "ARCHITECTURE-DECISION.md",
}


# ===========================================================================
# Extractors
# ===========================================================================

def _clean_text(text: str) -> str:
    """Normalize whitespace, strip null bytes, collapse blank lines."""
    text = text.replace("\x00", "")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


# ---------------------------------------------------------------------------
# 1. Cross-Risk-Overview.html extractor
# ---------------------------------------------------------------------------
class CrossRiskOverviewExtractor:
    """
    Extract text content from the Cross-Risk-Overview.html React SPA.

    Strategy:
    1. Parse the raw HTML with BeautifulSoup
    2. Extract text from script tags (React SPAs bundle content in JS)
    3. Extract any visible text in the HTML body
    4. Group extracted text by program section using keyword matching
    """

    PROGRAM_SECTIONS = {
        "Anti-Corruption": [
            "anti-corruption", "anti corruption", "tpi", "third-party intermediary",
            "third party intermediary", "tpdd", "due diligence", "gift", "ledger",
            "gl monitoring", "gl screening", "vendor",
        ],
        "Privacy": [
            "privacy", "onetrust", "one trust", "cookie", "dpia", "pia",
            "data mapping", "article 30", "consent", "dsr", "data subject",
            "mobile app", "data guidance", "gdpr",
        ],
        "Regulatory Compliance": [
            "regulatory", "compliance", "bridger", "lexisnexis", "screening",
            "sanctions", "kharon", "orbis", "case operator", "reg coe",
            "reg compliance", "mlex", "transparint",
        ],
        "Optimize": [
            "optimize", "tp inventory", "third party inventory", "3pi",
            "transfer pricing", "kpi", "data source", "persona",
        ],
        "AuditBoard": [
            "auditboard", "audit board", "opsaudit", "soxhub", "riskoversight",
            "crosscomply", "tcf", "tax control",
        ],
    }

    def __init__(self):
        self.html_path: Optional[Path] = None
        self.raw_html: str = ""
        self.extracted_chunks: list[dict] = []  # {program, text, source_label}

    def find_file(self) -> bool:
        for candidate in CROSS_RISK_OVERVIEW_CANDIDATES:
            if candidate.exists():
                self.html_path = candidate
                log.info(f"Found Cross-Risk-Overview.html at: {candidate}")
                return True
        log.warning("Cross-Risk-Overview.html NOT FOUND in any search path")
        return False

    def extract(self) -> list[dict]:
        """Main extraction — returns list of {program, text, source_label}."""
        if not self.html_path:
            return []

        try:
            from bs4 import BeautifulSoup
        except ImportError:
            log.error("beautifulsoup4 not installed — cannot parse HTML")
            return []

        self.raw_html = self.html_path.read_text(encoding="utf-8", errors="replace")
        soup = BeautifulSoup(self.raw_html, "html.parser")

        # Strategy 1: Extract text from the HTML body (SSR content, visible text)
        body_text = self._extract_body_text(soup)

        # Strategy 2: Extract text from script tags (React bundles inline content)
        script_text = self._extract_script_text(soup)

        # Strategy 3: Extract from data attributes (some React SPAs store content there)
        data_text = self._extract_data_attributes(soup)

        # Combine all extracted text
        all_text = body_text + "\n\n" + script_text + "\n\n" + data_text
        all_text = _clean_text(all_text)

        if len(all_text) < 100:
            log.warning(
                f"Cross-Risk-Overview.html yielded very little text ({len(all_text)} chars). "
                "This is likely a JS-only SPA that needs Playwright for full extraction."
            )

        # Split into chunks and classify by program
        self._classify_text(all_text)

        log.info(
            f"Cross-Risk-Overview: extracted {len(self.extracted_chunks)} chunks, "
            f"total {sum(len(c['text']) for c in self.extracted_chunks)} chars"
        )
        return self.extracted_chunks

    def _extract_body_text(self, soup) -> str:
        """Extract visible text from HTML body, excluding script/style tags."""
        body = soup.find("body")
        if not body:
            return ""

        # Remove script and style elements
        for tag in body.find_all(["script", "style", "noscript"]):
            tag.decompose()

        # Get text with newlines between block elements
        texts = []
        for element in body.find_all(
            ["h1", "h2", "h3", "h4", "h5", "h6", "p", "li", "td", "th",
             "div", "span", "section", "article", "figcaption", "blockquote"]
        ):
            text = element.get_text(separator=" ", strip=True)
            if text and len(text) > 3:
                texts.append(text)

        return "\n".join(texts)

    def _extract_script_text(self, soup) -> str:
        """
        Extract readable text from inline script tags.
        React SPAs often have content strings in the JS bundle.
        """
        texts = []
        for script in soup.find_all("script"):
            content = script.string or ""
            if not content or len(content) < 50:
                continue

            # Extract quoted strings that look like content (not code)
            # Match strings longer than 20 chars that contain spaces (likely prose)
            strings = re.findall(r'["\']([^"\']{20,})["\']', content)
            for s in strings:
                # Filter out things that look like code, URLs, CSS
                if any(skip in s for skip in [
                    "function", "return", "const ", "let ", "var ",
                    "http://", "https://", ".css", ".js", "webpack",
                    "createElement", "className", "onClick", "useState",
                    "__webpack", "module.exports", "import ", "export ",
                ]):
                    continue
                # Keep strings that look like natural language
                if " " in s and len(s.split()) >= 4:
                    texts.append(s)

            # Also try to extract JSX text content patterns
            # e.g., children: "Some text here"
            jsx_texts = re.findall(r'children:\s*["\']([^"\']{10,})["\']', content)
            texts.extend(jsx_texts)

            # Extract data from JSON-like structures in scripts
            json_texts = re.findall(r'"(?:title|text|content|description|label|heading)":\s*"([^"]{10,})"', content)
            texts.extend(json_texts)

        return "\n".join(texts)

    def _extract_data_attributes(self, soup) -> str:
        """Extract text from data-* attributes (some frameworks store content there)."""
        texts = []
        for tag in soup.find_all(attrs={"data-content": True}):
            texts.append(tag["data-content"])
        for tag in soup.find_all(attrs={"data-text": True}):
            texts.append(tag["data-text"])
        for tag in soup.find_all(attrs={"data-code": True}):
            # Mermaid diagrams might have content in data-code
            code = tag["data-code"]
            if len(code) > 20:
                texts.append(f"[Diagram] {code}")
        return "\n".join(texts)

    def _classify_text(self, all_text: str):
        """Split text into paragraphs and classify each by program."""
        # Split into meaningful paragraphs
        paragraphs = re.split(r"\n\n+", all_text)

        for para in paragraphs:
            para = para.strip()
            if len(para) < 20:
                continue

            para_lower = para.lower()
            matched_program = None
            best_score = 0

            for program, keywords in self.PROGRAM_SECTIONS.items():
                score = sum(1 for kw in keywords if kw in para_lower)
                if score > best_score:
                    best_score = score
                    matched_program = program

            if matched_program and best_score > 0:
                self.extracted_chunks.append({
                    "program": matched_program,
                    "text": para,
                    "source_label": f"Cross-Risk-Overview / {matched_program}",
                })
            elif len(para) > 50:
                # Unclassified but substantial text — keep as generic
                self.extracted_chunks.append({
                    "program": "General",
                    "text": para,
                    "source_label": "Cross-Risk-Overview / General",
                })


# ---------------------------------------------------------------------------
# 2. DOCX extractor
# ---------------------------------------------------------------------------
def extract_docx(file_path: Path) -> str:
    """Extract text from a .docx file using python-docx."""
    try:
        from docx import Document
    except ImportError:
        log.error("python-docx not installed — cannot read .docx files")
        return ""

    if not file_path.exists():
        log.warning(f"DOCX not found: {file_path}")
        return ""

    try:
        doc = Document(str(file_path))
        parts: list[str] = []

        # Extract paragraphs with heading markers
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style_name = (para.style.name or "").lower()
            if "heading" in style_name:
                level = "".join(c for c in style_name if c.isdigit()) or "1"
                parts.append(f"\n{'#' * int(level)} {text}\n")
            else:
                parts.append(text)

        # Extract tables
        for table_idx, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            if rows:
                parts.append(f"\n[Table {table_idx + 1}]\n" + "\n".join(rows))

        content = "\n".join(parts)
        log.info(f"DOCX extracted: {file_path.name} ({len(content)} chars)")
        return _clean_text(content)

    except Exception as e:
        log.error(f"Failed to extract DOCX {file_path.name}: {e}")
        return ""


# ---------------------------------------------------------------------------
# 3. Markdown extractor
# ---------------------------------------------------------------------------
def extract_markdown(file_path: Path) -> str:
    """Read a markdown file directly."""
    if not file_path.exists():
        log.warning(f"Markdown not found: {file_path}")
        return ""

    try:
        content = file_path.read_text(encoding="utf-8", errors="replace")
        log.info(f"MD extracted: {file_path.name} ({len(content)} chars)")
        return _clean_text(content)
    except Exception as e:
        log.error(f"Failed to read MD {file_path.name}: {e}")
        return ""


# ---------------------------------------------------------------------------
# 4. HTML extractor (for non-SPA HTML files)
# ---------------------------------------------------------------------------
def extract_html(file_path: Path) -> str:
    """Extract text from a standard HTML file using BeautifulSoup."""
    if not file_path.exists():
        log.warning(f"HTML not found: {file_path}")
        return ""

    try:
        from bs4 import BeautifulSoup
    except ImportError:
        log.error("beautifulsoup4 not installed")
        return ""

    try:
        raw = file_path.read_text(encoding="utf-8", errors="replace")
        soup = BeautifulSoup(raw, "html.parser")

        # Remove script and style
        for tag in soup.find_all(["script", "style", "noscript"]):
            tag.decompose()

        text = soup.get_text(separator="\n", strip=True)
        log.info(f"HTML extracted: {file_path.name} ({len(text)} chars)")
        return _clean_text(text)
    except Exception as e:
        log.error(f"Failed to extract HTML {file_path.name}: {e}")
        return ""


# ---------------------------------------------------------------------------
# 5. PDF extractor
# ---------------------------------------------------------------------------
def extract_pdf(file_path: Path) -> str:
    """Extract text from a PDF using pdfplumber."""
    if not file_path.exists():
        log.warning(f"PDF not found: {file_path}")
        return ""

    try:
        import pdfplumber
    except ImportError:
        log.error("pdfplumber not installed — cannot read PDFs")
        return ""

    try:
        texts: list[str] = []
        with pdfplumber.open(str(file_path)) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                page_text = page.extract_text() or ""
                if page_text.strip():
                    texts.append(f"[Page {page_num}]\n{page_text}")

                # Also extract tables
                tables = page.extract_tables()
                for table_idx, table in enumerate(tables):
                    if table:
                        rows = []
                        for row in table:
                            cells = [str(c or "").strip() for c in row]
                            rows.append(" | ".join(cells))
                        texts.append(f"[Page {page_num} Table {table_idx + 1}]\n" + "\n".join(rows))

        content = "\n\n".join(texts)
        log.info(f"PDF extracted: {file_path.name} ({len(content)} chars)")
        return _clean_text(content)
    except Exception as e:
        log.error(f"Failed to extract PDF {file_path.name}: {e}")
        return ""


# ---------------------------------------------------------------------------
# 6. Architecture PNG extractor (metadata only — actual vision analysis
#    requires Claude and is deferred to article generation phase)
# ---------------------------------------------------------------------------
def describe_architecture_pngs(product_name: str, arch_dir: Path) -> list[EvidenceSource]:
    """
    For architecture PNGs, we record what diagrams exist and their file paths.
    Actual image interpretation will be done by Claude during article generation.
    """
    sources: list[EvidenceSource] = []
    view_names = {
        "architecture-usecase-view.png": "Use Case View",
        "architecture-logical-view.png": "Logical View",
        "architecture-process-view.png": "Process View",
        "architecture-development-view.png": "Development View",
        "architecture-physical-view.png": "Physical View",
    }

    if not arch_dir.exists():
        log.warning(f"Architecture dir not found: {arch_dir}")
        return sources

    for filename, view_label in view_names.items():
        png_path = arch_dir / filename
        if png_path.exists():
            sources.append(EvidenceSource(
                type="architecture",
                name=f"4+1 UML / {product_name} / {view_label}",
                content=f"[Architecture Diagram: {view_label} for {product_name}. "
                        f"PNG file at: {png_path}. "
                        f"Requires Claude vision analysis during article generation.]",
                source_path=str(png_path),
            ))
            log.info(f"  Found arch diagram: {product_name} / {view_label}")
        else:
            log.debug(f"  Missing arch diagram: {product_name} / {view_label}")

    return sources


# ===========================================================================
# PRD file resolution
# ===========================================================================
def resolve_prd_path(filename: str) -> Optional[Path]:
    """Find the actual path for a PRD file, checking special paths first, then search dirs."""
    # Check special paths first
    if filename in SPECIAL_PRD_PATHS:
        path = SPECIAL_PRD_PATHS[filename]
        if path.exists():
            return path
        log.debug(f"Special path not found for {filename}: {path}")

    # Search in standard directories
    for search_dir in PRD_SEARCH_DIRS:
        if not search_dir.exists():
            continue
        candidate = search_dir / filename
        if candidate.exists():
            return candidate
        # Also try with different dash/hyphen encoding
        for f in search_dir.iterdir():
            if f.is_file() and f.name.replace("\u2013", "-").replace("\u2014", "-") == filename.replace("\u2013", "-").replace("\u2014", "-"):
                return f

    # Broader search in OneDrive
    for search_dir in [ONEDRIVE_CROSS_RISK, ONEDRIVE_INITIATIVES]:
        if not search_dir.exists():
            continue
        # Walk up to 4 levels deep
        for root, dirs, files in os.walk(search_dir):
            depth = root.replace(str(search_dir), "").count(os.sep)
            if depth > 4:
                dirs.clear()
                continue
            for f in files:
                # Normalize dashes for comparison
                f_norm = f.replace("\u2013", "-").replace("\u2014", "-")
                fn_norm = filename.replace("\u2013", "-").replace("\u2014", "-")
                if f_norm == fn_norm:
                    return Path(root) / f

    log.warning(f"PRD not found: {filename}")
    return None


# ===========================================================================
# Main pipeline
# ===========================================================================
def extract_product_evidence(product_def: dict, overview_data: list[dict]) -> ProductEvidence:
    """Extract all evidence for a single product."""
    pe = ProductEvidence(
        product=product_def["product"],
        slug=product_def["slug"],
        team=product_def["team"],
        program=product_def["program"],
        project_id=product_def.get("project_id", ""),
    )

    log.info(f"\n{'='*60}")
    log.info(f"Extracting: {pe.product} ({pe.slug})")
    log.info(f"{'='*60}")

    # --- 1. Cross-Risk-Overview chunks ---
    overview_program = product_def.get("overview_program")
    keywords = product_def.get("overview_keywords", [])

    if overview_program and overview_data:
        matching_chunks = []
        for chunk in overview_data:
            if chunk["program"] != overview_program and chunk["program"] != "General":
                continue
            chunk_lower = chunk["text"].lower()
            if any(kw.lower() in chunk_lower for kw in keywords):
                matching_chunks.append(chunk)

        if matching_chunks:
            combined = "\n\n".join(c["text"] for c in matching_chunks)
            pe.sources.append(EvidenceSource(
                type="presentation",
                name=f"Cross-Risk-Overview / {overview_program}",
                content=combined,
                source_path=str(CROSS_RISK_OVERVIEW_CANDIDATES[0]),
            ))
            log.info(f"  Overview: {len(matching_chunks)} matching chunks ({len(combined)} chars)")
        else:
            log.info(f"  Overview: no matching chunks for keywords {keywords}")
            pe.errors.append(f"No matching overview content for keywords: {keywords}")

    # --- 2. PRDs ---
    for prd_filename, prd_type in product_def.get("prds", []):
        prd_path = resolve_prd_path(prd_filename)
        if not prd_path:
            pe.errors.append(f"PRD not found: {prd_filename}")
            continue

        if prd_type == "docx":
            content = extract_docx(prd_path)
        elif prd_type == "md":
            content = extract_markdown(prd_path)
        elif prd_type == "html":
            content = extract_html(prd_path)
        else:
            pe.errors.append(f"Unknown PRD type: {prd_type} for {prd_filename}")
            continue

        if content:
            pe.sources.append(EvidenceSource(
                type="prd",
                name=prd_filename,
                content=content,
                source_path=str(prd_path),
            ))
        else:
            pe.errors.append(f"PRD extracted but empty: {prd_filename}")

    # --- 3. Architecture diagrams (4+1 UML PNGs) ---
    for arch_subdir in product_def.get("architecture", []):
        arch_dir = ARCH_UML_DIR / arch_subdir
        arch_sources = describe_architecture_pngs(pe.product, arch_dir)
        pe.sources.extend(arch_sources)

    # --- 4. Architecture PDFs ---
    for pdf_name in product_def.get("arch_pdfs", []):
        pdf_path = ARCH_PDF_PATHS.get(pdf_name)
        if pdf_path and pdf_path.exists():
            content = extract_pdf(pdf_path)
            if content:
                pe.sources.append(EvidenceSource(
                    type="architecture",
                    name=pdf_name,
                    content=content,
                    source_path=str(pdf_path),
                ))
        else:
            pe.errors.append(f"Architecture PDF not found: {pdf_name}")

    # --- 5. Extra architecture (3PI V2 docs) ---
    if "extra_arch" in product_def:
        for label in product_def["extra_arch"]:
            if label == "3PI V2":
                for doc_label, doc_path in THREE_PI_V2_DOCS.items():
                    if doc_path.exists():
                        if doc_path.suffix == ".html":
                            content = extract_html(doc_path)
                        elif doc_path.suffix == ".md":
                            content = extract_markdown(doc_path)
                        else:
                            continue
                        if content:
                            pe.sources.append(EvidenceSource(
                                type="architecture",
                                name=doc_label,
                                content=content,
                                source_path=str(doc_path),
                            ))
                    else:
                        pe.errors.append(f"3PI V2 doc not found: {doc_label}")

    # --- Compute tier and hash ---
    pe.compute_tier()
    pe.compute_hash()

    log.info(
        f"  Result: {pe.source_count} sources, tier={pe.evidence_tier}, "
        f"{sum(s.char_count for s in pe.sources)} total chars"
    )
    if pe.errors:
        log.warning(f"  Errors: {pe.errors}")

    return pe


def write_product_json(pe: ProductEvidence, output_dir: Path):
    """Write a single product evidence JSON."""
    data = {
        "product": pe.product,
        "slug": pe.slug,
        "team": pe.team,
        "program": pe.program,
        "project_id": pe.project_id,
        "sources": [
            {
                "type": s.type,
                "name": s.name,
                "content": s.content,
                "source_path": s.source_path,
                "char_count": s.char_count,
            }
            for s in pe.sources
        ],
        "evidence_tier": pe.evidence_tier,
        "source_count": pe.source_count,
        "source_hash": pe.source_hash,
    }

    out_path = output_dir / f"{pe.slug}.json"
    out_path.write_text(json.dumps(data, indent=2, ensure_ascii=False), encoding="utf-8")
    os.chmod(str(out_path), 0o600)
    log.info(f"  Written: {out_path}")


def write_manifest(results: list[dict], output_dir: Path):
    """Write extraction_manifest.json summarizing the pipeline run."""
    manifest = {
        "pipeline": "extract_product_evidence",
        "timestamp": __import__("datetime").datetime.now().isoformat(),
        "output_dir": str(output_dir),
        "total_products": len(results),
        "summary": {
            "rich": sum(1 for r in results if r["evidence_tier"] == "rich"),
            "standard": sum(1 for r in results if r["evidence_tier"] == "standard"),
            "thin": sum(1 for r in results if r["evidence_tier"] == "thin"),
            "failed": sum(1 for r in results if r.get("status") == "failed"),
        },
        "products": results,
    }

    manifest_path = output_dir / "extraction_manifest.json"
    manifest_path.write_text(json.dumps(manifest, indent=2, ensure_ascii=False), encoding="utf-8")
    log.info(f"\nManifest written: {manifest_path}")


def main():
    log.info("=" * 70)
    log.info("Product Evidence Extraction Pipeline")
    log.info("=" * 70)

    # --- Create output directory ---
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    os.chmod(str(OUTPUT_DIR), 0o700)  # Restrict permissions (internal data)
    log.info(f"Output directory: {OUTPUT_DIR}")

    # --- Step 1: Extract Cross-Risk-Overview ---
    log.info("\n--- Phase 1: Cross-Risk-Overview.html ---")
    overview_extractor = CrossRiskOverviewExtractor()
    overview_found = overview_extractor.find_file()
    overview_data: list[dict] = []
    if overview_found:
        overview_data = overview_extractor.extract()
        if not overview_data:
            log.warning(
                "Cross-Risk-Overview.html yielded no text. "
                "This is a React SPA — consider running with Playwright for full extraction. "
                "Continuing with PRD/architecture sources only."
            )
    else:
        log.warning(
            "Cross-Risk-Overview.html not found. Searched:\n"
            + "\n".join(f"  - {p}" for p in CROSS_RISK_OVERVIEW_CANDIDATES)
        )

    # --- Step 2: Process each product ---
    log.info("\n--- Phase 2: Per-product extraction ---")
    manifest_results: list[dict] = []

    for product_def in PRODUCTS:
        try:
            pe = extract_product_evidence(product_def, overview_data)
            write_product_json(pe, OUTPUT_DIR)

            manifest_results.append({
                "product": pe.product,
                "slug": pe.slug,
                "team": pe.team,
                "evidence_tier": pe.evidence_tier,
                "source_count": pe.source_count,
                "total_chars": sum(s.char_count for s in pe.sources),
                "source_hash": pe.source_hash,
                "source_types": list({s.type for s in pe.sources}),
                "errors": pe.errors,
                "status": "ok" if pe.sources else "no_sources",
            })

        except Exception as e:
            log.error(f"FAILED: {product_def['product']}: {e}")
            log.error(traceback.format_exc())
            manifest_results.append({
                "product": product_def["product"],
                "slug": product_def["slug"],
                "team": product_def["team"],
                "evidence_tier": "thin",
                "source_count": 0,
                "total_chars": 0,
                "source_hash": "",
                "source_types": [],
                "errors": [str(e)],
                "status": "failed",
            })

    # --- Step 3: Write manifest ---
    write_manifest(manifest_results, OUTPUT_DIR)

    # --- Step 4: Summary ---
    log.info("\n" + "=" * 70)
    log.info("EXTRACTION SUMMARY")
    log.info("=" * 70)

    tiers = {"rich": 0, "standard": 0, "thin": 0}
    no_sources = 0
    failed = 0
    for r in manifest_results:
        if r["status"] == "failed":
            failed += 1
        elif r["source_count"] == 0:
            no_sources += 1
        else:
            tiers[r["evidence_tier"]] += 1

    total_chars = sum(r["total_chars"] for r in manifest_results)

    log.info(f"  Total products: {len(manifest_results)}")
    log.info(f"  Rich (PRD+arch): {tiers['rich']}")
    log.info(f"  Standard (PRD or overview): {tiers['standard']}")
    log.info(f"  Thin (overview only): {tiers['thin']}")
    log.info(f"  No sources: {no_sources}")
    log.info(f"  Failed: {failed}")
    log.info(f"  Total extracted chars: {total_chars:,}")
    log.info(f"\n  Output: {OUTPUT_DIR}")
    log.info(f"  Manifest: {OUTPUT_DIR / 'extraction_manifest.json'}")

    # List products needing internet research
    internet_products = [r for r in manifest_results if r["source_count"] == 0 and r["status"] != "failed"]
    if internet_products:
        log.info(f"\n  Products needing internet research ({len(internet_products)}):")
        for r in internet_products:
            log.info(f"    - {r['product']}")

    # List products with errors
    error_products = [r for r in manifest_results if r["errors"]]
    if error_products:
        log.info(f"\n  Products with errors ({len(error_products)}):")
        for r in error_products:
            log.info(f"    - {r['product']}: {r['errors']}")

    log.info("\nDone.")
    return 0 if failed == 0 else 1


if __name__ == "__main__":
    sys.exit(main())
